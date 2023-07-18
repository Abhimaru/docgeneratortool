import docx as aw
from docx.oxml.xmlchemy import OxmlElement
from docx.oxml.ns import qn
import os
from aims import aims
from PIL import Image, ImageOps
from docx2pdf import convert

# Define Doc Details
heading = ''
headingSize = 20
aimSize = 16
paraSize = 12
fontOne = "Calibri"
fontTwo = "Consolas"
image_border_size = 5
id_no = "21CP315 - Abhishek Maru"
subject_name = "ANDROID LAB"

# Create a new document
doc = aw.Document()

# Adding Page Numbers


def create_element(name):
    return OxmlElement(name)


def create_attribute(element, name, value):
    element.set(qn(name), value)


def add_page_number(run):
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')

    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"

    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


# creating Header
def makeHeader(folder):

    txt = input("Enter Middle text for Heading (press Enter for Default): ")
    if (txt == ""):
        subject_name = "ANDROID"
    else:
        subject_name = txt

    txt = input("Enter Ending text for Heading (press Enter for Default): ")
    if (txt == ""):
        id_no = "21CP315 - Abhishek Maru"
    else:
        id_no = txt

    length = len(folder)
    i = 0
    for section in doc.sections:
        section.is_linked_to_previous = False
        if (i == length):
            return
        # print(folder[i])
        header = section.header
        paragraph = header.paragraphs[0]
        paragraph.text = f'{folder[i]}\t{subject_name}\t{id_no}'
        paragraph.style.font.name = fontOne
        paragraph.style.font.size = aw.shared.Pt(12)
        i += 1

# make output folder in sub folders in this directory if not exist


def makeFolder(myFolder):
    for folder in os.listdir(myFolder):
        if not os.path.exists(myFolder+"/"+folder+"/output"):
            os.mkdir(myFolder+"/"+folder+"/output")

# function for make image border


def add_image_border(input_image, output_image, border):
    img = Image.open(input_image)
    img_copy = img.copy()
    if isinstance(border, int) or isinstance(border, tuple):
        bimg = ImageOps.expand(img_copy, border=border)
    else:
        raise RuntimeError('Border is not an image or tuple')
    bimg.save(output_image)


# making border around page
def setPageBorder():
    for section in doc.sections:
        sec_pr = section._sectPr  # get the section properties el
        # create new borders el
        pg_borders = OxmlElement('w:pgBorders')
        # specifies how the relative positioning of the borders should be calculated
        pg_borders.set(qn('w:offsetFrom'), 'page')
        for border_name in ('top', 'left', 'bottom', 'right',):  # set all borders
            border_el = OxmlElement(f'w:{border_name}')
            border_el.set(qn('w:val'), 'single')  # a single line
            # for meaning of  remaining attrs please look docs
            border_el.set(qn('w:sz'), '1')
            border_el.set(qn('w:space'), '24')
            border_el.set(qn('w:color'), 'auto')
            pg_borders.append(border_el)  # register single border to border el
        sec_pr.append(pg_borders)  # apply border changes to section


# Set Document Margins to normal
sections = doc.sections
for section in sections:
    section.top_margin = aw.shared.Inches(1)
    section.bottom_margin = aw.shared.Inches(1)
    section.left_margin = aw.shared.Inches(1)
    section.right_margin = aw.shared.Inches(1)
    # Making A4 document
    section.page_height = aw.shared.Mm(297)
    section.page_width = aw.shared.Mm(210)
    section.header.is_linked_to_previous = True
    setPageBorder()

# Generating PDF


def generatePDF(file):
    convert(f'{file}.docx', f'{file}.pdf')


def makedoc(heading, aim, headingSize, aimSize, paraSize, folderName, filesNames, isLastPage=False):
    '''
    HEADINGS
    '''

    pg1 = doc.add_heading()
    pg1.underline = aw.shared.RGBColor(0, 0, 0)
    pg1.alignment = aw.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    pg1.style = doc.styles['Title']
    pg1 = pg1.add_run(heading)
    pg1.font.name = fontOne
    pg1.font.bold = True
    pg1.font.size = aw.shared.Pt(headingSize)
    pg1.font.color.rgb = aw.shared.RGBColor(0, 0, 0)

    '''
    AIM
    '''
    pg2 = doc.add_heading()
    pg2.alignment = aw.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
    pg2.paragraph_format.space_after = aw.shared.Pt(5)
    pg2 = pg2.add_run(f'Aim: {aim}')
    pg2.font.name = fontOne
    pg2.font.color.rgb = aw.shared.RGBColor(0, 0, 0)
    pg2.font.bold = True
    pg2.font.size = aw.shared.Pt(aimSize)
    pg2.font.color.rgb = aw.shared.RGBColor(0, 0, 0)

    # Read File and add to document
    for fileName in filesNames:
        pg3 = doc.add_paragraph()
        pg3 = pg3.add_run(fileName)
        pg3.font.underline = True
        pg3.font.bold = True
        pg3.font.name = fontOne
        pg3.font.size = aw.shared.Pt(aimSize-2)
        with open(f'{folderName}/{fileName}', 'r') as f:
            code = f.read()
        pg4 = doc.add_paragraph()
        pg4.paragraph_format.space_after = aw.shared.Pt(15)
        pg4 = pg4.add_run(code)
        pg4.alignment = aw.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        pg4.font.name = fontTwo
        pg4.font.size = aw.shared.Pt(paraSize)
        pg4.font.color.rgb = aw.shared.RGBColor(0, 0, 0)
        pg4.font.bold = False

    # Add Page Break
    doc.add_page_break()
    pg5 = doc.add_paragraph().add_run("OUTPUT:")
    pg5.alignment = aw.enum.text.WD_ALIGN_PARAGRAPH.LEFT
    pg5.font.name = fontOne
    pg5.font.bold = True
    pg5.font.size = aw.shared.Pt(aimSize)
    pg5.font.color.rgb = aw.shared.RGBColor(0, 0, 0)

    # Add Screenshot to document
    tmp = folderName+'/output/'
    for image in os.listdir(tmp):
        add_image_border(tmp+image, f'{tmp}tmp_{image}', image_border_size)
        doc.add_picture(f'{tmp}tmp_{image}', width=aw.shared.Inches(3.13))
        os.remove(f'{tmp}tmp_{image}')

    if not isLastPage:
        doc.add_page_break()
        # add new section
        doc.add_section().header.is_linked_to_previous = False


def makedocIndividual(heading, aim, headingSize, aimSize, paraSize, image, file, isLastPage=False):
    '''
    HEADINGS
    '''
    pg1 = doc.add_heading()
    pg1.underline = aw.shared   .RGBColor(0, 0, 0)
    pg1.alignment = aw.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    pg1.style = doc.styles['Title']
    pg1 = pg1.add_run(heading)
    pg1.font.name = fontOne
    pg1.font.bold = True
    pg1.font.size = aw.shared.Pt(headingSize)
    pg1.font.color.rgb = aw.shared.RGBColor(0, 0, 0)

    '''
    AIM
    '''
    pg2 = doc.add_heading()
    pg2.alignment = aw.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
    pg2 = pg2.add_run(f'Aim: {aim}')
    pg2.font.name = fontOne
    pg2.font.color.rgb = aw.shared.RGBColor(0, 0, 0)
    pg2.font.bold = True
    pg2.font.size = aw.shared.Pt(aimSize)
    pg2.font.color.rgb = aw.shared.RGBColor(0, 0, 0)

    # Read File and add to document
    pg3 = doc.add_paragraph()
    pg3.paragraph_format.space_before = aw.shared.Pt(15)
    pg3 = pg3.add_run("CODE:")
    pg3.alignment = aw.enum.text.WD_ALIGN_PARAGRAPH.LEFT
    pg3.font.name = fontOne
    pg3.font.bold = True
    pg3.font.underline = True
    pg3.font.size = aw.shared.Pt(aimSize-2)
    pg3.font.color.rgb = aw.shared.RGBColor(0, 0, 0)

    with open(file, 'r') as f:
        code = f.read()
    pg4 = doc.add_paragraph()
    pg4.paragraph_format.space_after = aw.shared.Pt(15)
    pg4 = pg4.add_run(code)
    pg4.alignment = aw.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    pg4.font.name = fontTwo
    pg4.font.size = aw.shared.Pt(paraSize)
    pg4.font.color.rgb = aw.shared.RGBColor(0, 0, 0)
    pg4.font.bold = False

    # Add Page Break
    pg5 = doc.add_paragraph().add_run("OUTPUT:")
    pg5.alignment = aw.enum.text.WD_ALIGN_PARAGRAPH.LEFT
    pg5.font.name = fontOne
    pg5.font.bold = True
    pg5.font.underline = True
    pg5.font.size = aw.shared.Pt(aimSize)
    pg5.font.color.rgb = aw.shared.RGBColor(0, 0, 0)

    # Add Screenshot to document
    tmp = 'outputs/'
    add_image_border(tmp+image, f'{tmp}tmp_{image}', image_border_size)
    doc.add_picture(f'{tmp}tmp_{image}', width=aw.shared.Inches(3.13))
    os.remove(f'{tmp}tmp_{image}')

    if not isLastPage:
        doc.add_page_break()
        # add new section
        doc.add_section().header.is_linked_to_previous = False


java_files = []
xml_files = []
other_files = []
filesNames = []
folderNames = []
parentFolderName = "LABS"

opt = input(
    "1. For individual files\n2. For individual folders\nChoose options (1/2):")
if opt == '1':
    os.chdir("LABS")
    if not os.path.exists("outputs"):
        os.mkdir("outputs")

    def sort_file(file_list):
        return sorted(file_list, key=lambda x: int(x.split('-')[0]))

    def sort_img(img_list):
        return sorted(img_list, key=lambda x: int(x.split('.')[0]))

    file_list = []
    img_list = []
    i = 1
    for file in os.listdir():
        if os.path.isfile(file):
            file_list.append(file)
        mystr = f'{i}.png'
        img_list.append(mystr)
        i += 1

    img_list = sort_img(img_list)
    file_list = sort_file(file_list)
    count = 1
    heading_list = []
    tmp = 0
    for file in file_list:
        try:
            if not tmp == len(file_list):
                img = img_list[tmp]

            heading = f'EXPERIMENT-{count}'
            heading_list.append(heading)
            isLastLab = False
            aim = aims[heading] if aims.keys().__contains__(heading) else ""
            if file == file_list[-1]:
                isLastLab = True
            makedocIndividual(heading, aim, headingSize, aimSize,
                              paraSize, img, file, isLastLab)
            count += 1
            tmp += 1
        except:
            print("No image found")
    img_list.clear()
    file_list.clear()
    txt = input("do you want to add Header? (y/n)")
    if (txt == 'y' or txt == 'Y'):
        makeHeader(heading_list)

elif opt == '2':
    parentFolderName = "LABS"
    makeFolder(parentFolderName)

    for folder in os.listdir(parentFolderName):
        # If folder is directory then append it to folderNames
        if os.path.isdir(os.path.join(parentFolderName, folder)):
            folderNames.append(folder)

    # Set Current Directory to parentFolderName
    os.chdir(parentFolderName)
    isLastLab = False
    op = input("1. For Single LAB\n2. For All LAB\nChoose options (1/2):")
    if op == '1':
        print("Lab No\tLab Name")
        for i in range((len(folderNames))):
            print(f'{i+1}\t{folderNames[i]}')
        labNo = int(
            input("Enter Lab No: (Any Number from above list,Otherwise it will be exit) : "))
        if labNo > len(folderNames) or labNo < 1:
            print("Invalid Input")
            exit()
        else:
            isLastLab = True
            folderName = folderNames[labNo-1]
            heading = folderName
            aim = aims[folderName] if aims[folderName] != None else ""
            for file in os.listdir(folderName):
                file_name = os.path.join(folderName, file)
                if os.path.isfile(file_name):
                    if file.endswith('.java') and file != 'MainActivity.java':
                        java_files.append(file)
                    elif file.endswith('.java') and file == 'MainActivity.java':
                        java_files.insert(0, file)
                    elif file.endswith('.xml') and file != 'activity_main.xml':
                        xml_files.append(file)
                    elif file.endswith('.xml') and file == 'activity_main.xml':
                        xml_files.insert(0, file)
                    else:
                        other_files.append(file)

                filesNames = java_files + xml_files + other_files

            # call makedoc function
            makedoc(heading, aim, headingSize, aimSize,
                    paraSize, folderName, filesNames, isLastLab)
            java_files.clear()
            xml_files.clear()
            other_files.clear()
            filesNames.clear()

        txt = input("do you want to add Header? (y/n)")
        if (txt == 'y' or txt == 'Y'):
            makeHeader(folderNames)

    elif op == '2':
        def sort_folder(file_list):
            return sorted(file_list, key=lambda x: int(x.split('-')[1]))

        folderNames = sort_folder(folderNames)
        for folderName in folderNames:
            heading = folderName
            aim = aims[folderName] if aims.keys(
            ).__contains__(folderName) else ""
            for file in os.listdir(folderName):
                file_name = os.path.join(folderName, file)
                if os.path.isfile(file_name):
                    if file.endswith('.java') and file != 'MainActivity.java':
                        java_files.append(file)
                    elif file.endswith('.java') and file == 'MainActivity.java':
                        java_files.insert(0, file)
                    elif file.endswith('.xml') and file != 'activity_main.xml':
                        xml_files.append(file)
                    elif file.endswith('.xml') and file == 'activity_main.xml':
                        xml_files.insert(0, file)
                    else:
                        other_files.append(file)

                filesNames = java_files + xml_files + other_files

            if folderName == folderNames[-1]:
                isLastLab = True
            # print(f'Folder Name: {folderName}')
            # print(f'Files: {filesNames}')
            # call makedoc function
            makedoc(heading, aim, headingSize, aimSize,
                    paraSize, folderName, filesNames, isLastLab)
            isLastLab = False
            java_files.clear()
            xml_files.clear()
            other_files.clear()
            filesNames.clear()
        txt = input("do you want to add Header? (y/n)")
        if (txt == 'y' or txt == 'Y'):
            makeHeader(folderNames)

    else:
        print("Invalid Input")
        exit()
else:
    print("Invalid Input")
    exit()


txt = input("do you want to add Page Number in Footer? (y/n)")
if (txt == 'y' or txt == 'Y'):
    footer = doc.sections[0].footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = aw.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
    f_run = footer_para.add_run()
    add_page_number(f_run)

# Save document
nameOftheFile = input("Enter the name of the file to save: ")
os.chdir('..')
doc.save(f'{nameOftheFile}.docx')

txt = input("Do you want to generate PDF? (y/n): ")
if txt == 'y' or txt == 'Y':
    generatePDF(nameOftheFile)
    os.system(f'start {nameOftheFile}.pdf')

txt = input("Do you want to Open Word File? (y/n): ")
if txt == 'y' or txt == 'Y':
    os.system(f'start {nameOftheFile}.docx')

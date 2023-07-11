import docx as aw
from docx.oxml.xmlchemy import OxmlElement
from docx.oxml.ns import qn
import os
from aims import aims

# Define Doc Details
heading = ''
headingSize = 20
aimSize = 16
paraSize = 12
fontOne = "Calibri"
fontTwo = "Consolas"

# Create a new document
doc = aw.Document()

# make output folder in sub folders in this directory if not exist


def makeFolder(myFolder):
    for folder in os.listdir(myFolder):
        if not os.path.exists(myFolder+"/"+folder+"/output"):
            os.mkdir(myFolder+"/"+folder+"/output")


def setPageBorder():
    for section in doc.sections:
        sec_pr = section._sectPr  # get the section properties el
        # create new borders el
        pg_borders = OxmlElement('w:pgBorders')
        # specifies how the relative positioning of the borders should be calculated
        pg_borders.set(qn('w:offsetFrom'), 'page')
        for border_name in ('top', 'left', 'bottom', 'right',):  # set all borders
            border_el = OxmlElement(f'w:{border_name}')
            border_el.set(qn('w:val'), 'single')  # a single line
            # for meaning of  remaining attrs please look docs
            border_el.set(qn('w:sz'), '1')
            border_el.set(qn('w:space'), '24')
            border_el.set(qn('w:color'), 'auto')
            pg_borders.append(border_el)  # register single border to border el
        sec_pr.append(pg_borders)  # apply border changes to section


# Set Document Margins to normal
sections = doc.sections
for section in sections:
    section.top_margin = aw.shared.Inches(1)
    section.bottom_margin = aw.shared.Inches(1)
    section.left_margin = aw.shared.Inches(1)
    section.right_margin = aw.shared.Inches(1)
    # Making A4 document
    section.page_height = aw.shared.Mm(297)
    section.page_width = aw.shared.Mm(210)
    setPageBorder()


def makedoc(heading, aim, headingSize, aimSize, paraSize, folderName, filesNames):
    '''
    HEADINGS
    '''
    pg1 = doc.add_paragraph()
    pg1.alignment = aw.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    pg1 = pg1.add_run(heading)
    pg1.font.name = fontOne
    pg1.font.bold = True
    pg1.font.underline = True
    pg1.font.size = aw.shared.Pt(headingSize)
    pg1.font.color.rgb = aw.shared.RGBColor(0, 0, 0)

    '''
    AIM
    '''
    pg2 = doc.add_paragraph()
    pg2.alignment = aw.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
    pg2 = pg2.add_run(f'Aim: {aim}')
    pg2.font.name = fontOne
    pg2.font.color.rgb = aw.shared.RGBColor(0, 0, 0)
    pg2.font.bold = True
    pg2.font.size = aw.shared.Pt(aimSize)
    pg2.font.color.rgb = aw.shared.RGBColor(0, 0, 0)

    # Read File and add to document
    for fileName in filesNames:
        pg3 = doc.add_paragraph()
        pg3.style = 'List Bullet'
        pg3 = pg3.add_run(fileName)
        pg3.font.underline = True
        pg3.font.bold = True
        pg3.font.size = aw.shared.Pt(aimSize-2)
        with open(f'{folderName}/{fileName}', 'r') as f:
            code = f.read()
        pg4 = doc.add_paragraph()
        pg4.paragraph_format.space_after = aw.shared.Pt(15)
        pg4 = pg4.add_run(code)
        pg4.alignment = aw.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        pg4.font.name = fontTwo
        pg4.font.size = aw.shared.Pt(paraSize)
        pg4.font.color.rgb = aw.shared.RGBColor(0, 0, 0)
        pg4.font.bold = False

    # Add Page Break
    doc.add_page_break()
    pg5 = doc.add_paragraph().add_run("OUTPUT:")
    pg5.alignment = aw.enum.text.WD_ALIGN_PARAGRAPH.LEFT
    pg5.font.name = fontOne
    pg5.font.bold = True
    pg5.font.underline = True
    pg5.font.size = aw.shared.Pt(aimSize)
    pg5.font.color.rgb = aw.shared.RGBColor(0, 0, 0)

    # Add Screenshot to document
    tmp = folderName+'/output/'
    for image in os.listdir(tmp):
        image_name = os.path.basename(image)
        doc.add_picture(tmp+image_name, width=aw.shared.Inches(3.13))

    doc.add_page_break()


java_files = []
xml_files = []
other_files = []
filesNames = []
folderNames = []
parentFolderName = "LABS"

makeFolder(parentFolderName)

for folder in os.listdir(parentFolderName):
    # If folder is directory then append it to folderNames
    if os.path.isdir(os.path.join(parentFolderName, folder)):
        folderNames.append(folder)

# Set Current Directory to parentFolderName
os.chdir(parentFolderName)

op = input("1. For Single LAB\n2. For All LAB\nChoose options (1/2):")
if op == '1':
    print("Lab No\tLab Name")
    for i in range((len(folderNames))):
        print(f'{i+1}\t{folderNames[i]}')
    labNo = int(
        input("Enter Lab No: (Any Number from above list,Otherwise it will be exit) : "))
    if labNo > len(folderNames) or labNo < 1:
        print("Invalid Input")
        exit()
    else:
        folderName = folderNames[labNo-1]
        heading = folderName
        aim = aims[folderName] if aims[folderName] != None else ""
        for file in os.listdir(folderName):
            file_name = os.path.join(folderName, file)
            if os.path.isfile(file_name):
                if file.endswith('.java') and file != 'MainActivity.java':
                    java_files.append(file)
                elif file.endswith('.java') and file == 'MainActivity.java':
                    java_files.insert(0, file)
                elif file.endswith('.xml') and file != 'activity_main.xml':
                    xml_files.append(file)
                elif file.endswith('.xml') and file == 'activity_main.xml':
                    xml_files.insert(0, file)
                else:
                    other_files.append(file)

            filesNames = java_files + xml_files + other_files

        # call makedoc function
        makedoc(heading, aim, headingSize, aimSize,
                paraSize, folderName, filesNames)
        java_files.clear()
        xml_files.clear()
        other_files.clear()
        filesNames.clear()

elif op == '2':
    for folderName in folderNames:
        heading = folderName
        aim = aims[folderName] if aims.keys().__contains__(folderName) else ""
        for file in os.listdir(folderName):
            file_name = os.path.join(folderName, file)
            if os.path.isfile(file_name):
                if file.endswith('.java') and file != 'MainActivity.java':
                    java_files.append(file)
                elif file.endswith('.java') and file == 'MainActivity.java':
                    java_files.insert(0, file)
                elif file.endswith('.xml') and file != 'activity_main.xml':
                    xml_files.append(file)
                elif file.endswith('.xml') and file == 'activity_main.xml':
                    xml_files.insert(0, file)
                else:
                    other_files.append(file)

            filesNames = java_files + xml_files + other_files

        # print(f'Folder Name: {folderName}')
        # print(f'Files: {filesNames}')
        # call makedoc function
        makedoc(heading, aim, headingSize, aimSize,
                paraSize, folderName, filesNames)
        java_files.clear()
        xml_files.clear()
        other_files.clear()
        filesNames.clear()

else:
    print("Invalid Input")
    exit()

# Save document
nameOftheFile = input("Enter the name of the file: ")
os.chdir('..')
doc.save(f'{nameOftheFile}.docx')
# Open Document
os.system(f'start {nameOftheFile}.docx')
